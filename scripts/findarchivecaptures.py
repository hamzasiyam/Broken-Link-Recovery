import os
import json
import subprocess
import pandas as pd
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
from tkinter import Tk, filedialog, messagebox, Label, Entry, Button, StringVar, Toplevel, ttk

PROFILE_DIR = "./profiles"
REPORTS_DIR = "./reports"
PROCESSED_DIR = os.path.join(REPORTS_DIR, "processed")
RAW_FILES_DIR = os.path.join(REPORTS_DIR, "raw_files")
LOGOS_DIR = os.path.join(REPORTS_DIR, "all_logos")
SCRIPTS_DIR = "./scripts"
PROFILE_FILE = os.path.join(PROFILE_DIR, "snapshot_analysis_profiles.json")

# Ensure the necessary directories exist
os.makedirs(PROFILE_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)
os.makedirs(PROCESSED_DIR, exist_ok=True)
os.makedirs(RAW_FILES_DIR, exist_ok=True)
os.makedirs(LOGOS_DIR, exist_ok=True)
os.makedirs(SCRIPTS_DIR, exist_ok=True)

def get_snapshots(url):
    result = subprocess.run(['waybackpack', url, '--list'], capture_output=True, text=True)
    if result.returncode != 0:
        print("Error fetching snapshots")
        print(result.stderr)
        return []
    
    snapshots = result.stdout.strip().split('\n')
    return snapshots

def extract_date_from_link(link):
    timestamp = link.split('/')[4]
    date = datetime.strptime(timestamp, '%Y%m%d%H%M%S')
    return date

def split_date_components(date):
    return date.strftime('%B %d, %Y'), date.strftime('%I:%M:%S %p')

def set_font_color(run, hex_color):
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)

def create_summary_report(snapshots, domain, logo_path, logo_height_percent, title_color_hex, heading_color_hex):
    filename = os.path.join(PROCESSED_DIR, f'{domain}_summary_report.docx')
    data = [(i+1, *split_date_components(extract_date_from_link(snap))) for i, snap in enumerate(snapshots)]
    df = pd.DataFrame(data, columns=['Capture', 'Date', 'Time'])

    summary = {
        'Total Captures': len(snapshots),
        'First Capture': {
            'Date': df['Date'].iloc[0],
            'Time': df['Time'].iloc[0]
        },
        'Last Capture': {
            'Date': df['Date'].iloc[-1],
            'Time': df['Time'].iloc[-1]
        }
    }

    doc = Document()
    
    # Add the logo above the title
    if logo_path:
        from PIL import Image
        logo = Image.open(logo_path)
        logo_width, logo_height = logo.size
        logo_height = int(logo_height * int(logo_height_percent) / 100)
        logo_width = int(logo_width * logo_height / logo.size[1])
        doc.add_picture(logo_path, width=Inches(logo_width / 96), height=Inches(logo_height / 96))

    title_run = doc.add_heading(f'Snapshot Summary Report for {domain}', 0).runs[0]
    set_font_color(title_run, title_color_hex)

    heading_1_run = doc.add_heading('Introduction', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph('This report provides a summary of the snapshots captured from the specified URL. The snapshots represent historical captures of the website, which are the first step in our process of identifying and fixing broken links. By analyzing these snapshots, we can gain insights into the changes and updates made to the website over time.')

    heading_1_run = doc.add_heading('Methodology', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph('The data was collected to provide a timeline of the website\'s state at various points in time. These snapshots help us understand how the website has evolved and where issues such as broken links may have arisen.')

    heading_1_run = doc.add_heading('Summary of Captures', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph(f"Total Captures: {summary['Total Captures']}")
    
    heading_2_run = doc.add_heading('First Capture', level=2).runs[0]
    set_font_color(heading_2_run, heading_color_hex)
    doc.add_paragraph(f"Date: {summary['First Capture']['Date']}")
    doc.add_paragraph(f"Time: {summary['First Capture']['Time']}")
    
    heading_2_run = doc.add_heading('Last Capture', level=2).runs[0]
    set_font_color(heading_2_run, heading_color_hex)
    doc.add_paragraph(f"Date: {summary['Last Capture']['Date']}")
    doc.add_paragraph(f"Time: {summary['Last Capture']['Time']}")

    doc.save(filename)

def create_detailed_analysis_report(snapshots, domain, logo_path, logo_height_percent, title_color_hex, heading_color_hex, column_color_hex):
    filename = os.path.join(PROCESSED_DIR, f'{domain}_detailed_analysis_report.docx')
    data = [(i+1, *split_date_components(extract_date_from_link(snap))) for i, snap in enumerate(snapshots)]
    df = pd.DataFrame(data, columns=['Capture', 'Date', 'Time'])

    doc = Document()
    
    # Add the logo above the title
    if logo_path:
        from PIL import Image
        logo = Image.open(logo_path)
        logo_width, logo_height = logo.size
        logo_height = int(logo_height * int(logo_height_percent) / 100)
        logo_width = int(logo_width * logo_height / logo.size[1])
        doc.add_picture(logo_path, width=Inches(logo_width / 96), height=Inches(logo_height / 96))

    title_run = doc.add_heading(f'Detailed Snapshot Analysis for {domain}', 0).runs[0]
    set_font_color(title_run, title_color_hex)

    heading_1_run = doc.add_heading('Introduction', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph('This document provides a detailed analysis of the snapshots captured from the specified URL. Each snapshot represents a historical state of the website. This detailed analysis is a crucial step in our process of identifying and fixing broken links across the website by examining each snapshot and checking all the files and links within them.')

    heading_1_run = doc.add_heading('Detailed Analysis', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph(f"Total Captures: {len(snapshots)}")
    doc.add_paragraph('The table below lists the snapshots in ascending order, organized by date and time.')

    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Capture'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Time'

    for cell in hdr_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(12)
        shading_elm_1 = OxmlElement("w:shd")
        shading_elm_1.set(qn("w:fill"), column_color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm_1)

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Capture'])
        row_cells[1].text = row['Date']
        row_cells[2].text = row['Time']

    # Add borders to the table
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    heading_1_run = doc.add_heading('Next Steps', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)

    def add_bold_paragraph(doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        return p

    add_bold_paragraph(doc, '1. Examination of Snapshots:')
    doc.add_paragraph('Each snapshot will be analyzed to check for broken links. This includes going through every file and link within each historical capture.')

    add_bold_paragraph(doc, '2. Identification and Fixing of Broken Links:')
    doc.add_paragraph('Any broken links identified during the analysis will be fixed.')

    add_bold_paragraph(doc, '3. Follow-Up Reports:')
    doc.add_paragraph('Additional reports will be provided throughout the process to document the progress and any further findings.')

    doc.save(filename)

def browse_file(variable):
    file_path = filedialog.askopenfilename(initialdir=LOGOS_DIR, filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])
    variable.set(file_path)

def generate_reports():
    global url_var, logo_path_var, logo_height_percent_var, title_color_hex_var, heading_color_hex_var, column_color_hex_var
    url = url_var.get()
    if not url.startswith("http://") and not url.startswith("https://"):
        url = "http://" + url
    domain = url.split('//')[1].split('/')[0]
    snapshots = get_snapshots(url)
    logo_path = logo_path_var.get()
    logo_height_percent = logo_height_percent_var.get()
    title_color_hex = title_color_hex_var.get().lstrip('#')
    heading_color_hex = heading_color_hex_var.get().lstrip('#')
    column_color_hex = column_color_hex_var.get().lstrip('#')
    if snapshots:
        create_summary_report(snapshots, domain, logo_path, logo_height_percent, title_color_hex, heading_color_hex)
        create_detailed_analysis_report(snapshots, domain, logo_path, logo_height_percent, title_color_hex, heading_color_hex, column_color_hex)
        messagebox.showinfo("Success", f"Reports saved to {PROCESSED_DIR}")
    else:
        messagebox.showerror("Error", "No snapshots found or there was an error.")

def save_profile(name, profile_data):
    profiles = load_profiles()
    profiles[name] = profile_data

    with open(PROFILE_FILE, "w") as f:
        json.dump(profiles, f, indent=4)

def load_profiles():
    if not os.path.exists(PROFILE_FILE):
        # Create a default profile file if it doesn't exist
        default_profiles = {
            "default": {
                "url": "",
                "logo_path": "",
                "logo_height_percent": "50",
                "title_color_hex": "000000",
                "heading_color_hex": "000000",
                "column_color_hex": "FFFFFF"
            }
        }
        with open(PROFILE_FILE, "w") as f:
            json.dump(default_profiles, f, indent=4)
        return default_profiles
    else:
        with open(PROFILE_FILE, "r") as f:
            return json.load(f)

def select_profile(profile_name_var, vars):
    profiles = load_profiles()
    profile_name = profile_name_var.get()
    if profile_name in profiles:
        profile = profiles[profile_name]
        vars['url_var'].set(profile["url"])
        vars['logo_path_var'].set(profile["logo_path"])
        vars['logo_height_percent_var'].set(profile["logo_height_percent"])
        vars['title_color_hex_var'].set(profile["title_color_hex"])
        vars['heading_color_hex_var'].set(profile["heading_color_hex"])
        vars['column_color_hex_var'].set(profile["column_color_hex"])

def delete_profile(profile_name_var):
    profiles = load_profiles()
    profile_name = profile_name_var.get()
    if profile_name in profiles:
        del profiles[profile_name]
        with open(PROFILE_FILE, "w") as f:
            json.dump(profiles, f, indent=4)
        return True
    return False

def main():
    global url_var, logo_path_var, logo_height_percent_var, title_color_hex_var, heading_color_hex_var, column_color_hex_var

    def create_profile(existing_profile=None):
        profile_window = Toplevel(root)
        profile_window.title("Create/Edit Profile")
        profile_window.geometry("550x450")

        profile_name = StringVar(value=existing_profile if existing_profile else "")
        profile_url = StringVar()
        profile_logo_path = StringVar()
        profile_logo_height_percent = StringVar(value="50")
        profile_title_color_hex = StringVar(value="000000")
        profile_heading_color_hex = StringVar(value="000000")
        profile_column_color_hex = StringVar(value="FFFFFF")

        if existing_profile:
            profiles = load_profiles()
            profile = profiles[existing_profile]
            profile_url.set(profile["url"])
            profile_logo_path.set(profile["logo_path"])
            profile_logo_height_percent.set(profile["logo_height_percent"])
            profile_title_color_hex.set(profile["title_color_hex"])
            profile_heading_color_hex.set(profile["heading_color_hex"])
            profile_column_color_hex.set(profile["column_color_hex"])

        Label(profile_window, text="Profile Name").grid(row=0, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_name, width=50).grid(row=0, column=1)

        Label(profile_window, text="URL").grid(row=1, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_url, width=50).grid(row=1, column=1)

        Label(profile_window, text="Logo File").grid(row=2, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_path, width=50).grid(row=2, column=1)
        Button(profile_window, text="Browse", command=lambda: browse_file(profile_logo_path)).grid(row=2, column=2)

        Label(profile_window, text="Logo Height (%)").grid(row=3, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_height_percent, width=10).grid(row=3, column=1, sticky='w')

        Label(profile_window, text="Title Color (hex)").grid(row=4, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_title_color_hex, width=10).grid(row=4, column=1, sticky='w')

        Label(profile_window, text="Heading Color (hex)").grid(row=5, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_heading_color_hex, width=10).grid(row=5, column=1, sticky='w')

        Label(profile_window, text="Column Color (hex)").grid(row=6, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_column_color_hex, width=10).grid(row=6, column=1, sticky='w')

        def save_new_profile():
            profile_data = {
                "url": profile_url.get(),
                "logo_path": profile_logo_path.get(),
                "logo_height_percent": profile_logo_height_percent.get(),
                "title_color_hex": profile_title_color_hex.get(),
                "heading_color_hex": profile_heading_color_hex.get(),
                "column_color_hex": profile_column_color_hex.get()
            }
            save_profile(profile_name.get(), profile_data)
            profile_window.destroy()
            load_profile_names()

        Button(profile_window, text="Save Profile", command=save_new_profile).grid(row=7, columnspan=3, pady=20)

    def load_profile_names():
        profiles = load_profiles()
        profile_names = list(profiles.keys())
        profile_name_combobox['values'] = profile_names

    def select_profile_callback(event):
        select_profile(profile_name_var, {
            'url_var': url_var,
            'logo_path_var': logo_path_var,
            'logo_height_percent_var': logo_height_percent_var,
            'title_color_hex_var': title_color_hex_var,
            'heading_color_hex_var': heading_color_hex_var,
            'column_color_hex_var': column_color_hex_var
        })

    def delete_profile_callback():
        if delete_profile(profile_name_var):
            load_profile_names()
            messagebox.showinfo("Success", "Profile deleted successfully.")
        else:
            messagebox.showerror("Error", "Profile not found.")

    root = Tk()
    root.title("Snapshot Analysis Report Generator")
    root.geometry("600x600")

    url_var = StringVar()
    logo_path_var = StringVar()
    logo_height_percent_var = StringVar(value="50")
    title_color_hex_var = StringVar(value="000000")
    heading_color_hex_var = StringVar(value="000000")
    column_color_hex_var = StringVar(value="FFFFFF")

    # Summary of what the program does
    summary_label = Label(root, text="This tool generates snapshot analysis reports for websites. "
                                     "It retrieves historical snapshots from a URL and creates detailed "
                                     "reports that help identify and fix broken links.",
                          wraplength=500, justify="left", font=("Arial", 10, "italic"))
    summary_label.grid(row=0, columnspan=3, padx=10, pady=10)

    Label(root, text="Enter URL:").grid(row=1, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=url_var, width=50).grid(row=1, column=1, padx=10, pady=10, sticky='w')

    Label(root, text="Logo File:").grid(row=2, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=logo_path_var, width=50).grid(row=2, column=1, padx=10, pady=10, sticky='w')
    Button(root, text="Browse", command=lambda: browse_file(logo_path_var)).grid(row=2, column=2, padx=10, pady=10, sticky='w')

    Label(root, text="Logo Height (%):").grid(row=3, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=logo_height_percent_var, width=10).grid(row=3, column=1, padx=10, pady=10, sticky='w')

    Label(root, text="Title Font Color (Hex):").grid(row=4, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=title_color_hex_var, width=10).grid(row=4, column=1, padx=10, pady=10, sticky='w')

    Label(root, text="Heading Font Color (Hex):").grid(row=5, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=heading_color_hex_var, width=10).grid(row=5, column=1, padx=10, pady=10, sticky='w')

    Label(root, text="Column Shading Color (Hex):").grid(row=6, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=column_color_hex_var, width=10).grid(row=6, column=1, padx=10, pady=10, sticky='w')

    Button(root, text="Generate Reports", command=generate_reports).grid(row=7, columnspan=3, padx=10, pady=20)

    Label(root, text="Profile Name").grid(row=8, column=0, pady=10, sticky='w')
    profile_name_var = StringVar()
    profile_name_combobox = ttk.Combobox(root, textvariable=profile_name_var, state='readonly')
    profile_name_combobox.grid(row=8, column=1, pady=10, padx=10)
    profile_name_combobox.bind("<<ComboboxSelected>>", select_profile_callback)

    Button(root, text="Create Profile", command=lambda: create_profile()).grid(row=9, column=0, pady=10)
    Button(root, text="Edit Profile", command=lambda: create_profile(existing_profile=profile_name_var.get())).grid(row=9, column=1, pady=10)
    Button(root, text="Delete Profile", command=delete_profile_callback).grid(row=9, column=2, pady=10)

    load_profile_names()

    root.mainloop()

if __name__ == "__main__":
    main()
