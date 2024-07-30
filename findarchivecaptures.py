import os
import subprocess
import pandas as pd
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
from tkinter import Tk, filedialog, messagebox, Label, Entry, Button, StringVar

def get_snapshots(url):
    result = subprocess.run(['waybackpack', url, '--list'], capture_output=True, text=True)
    if result.returncode != 0:
        print("Error fetching snapshots")
        print(result.stderr)
        return []
    
    snapshots = result.stdout.strip().split('\n')
    return snapshots

def extract_date_from_link(link):
    timestamp = link.split('/')[4]
    date = datetime.strptime(timestamp, '%Y%m%d%H%M%S')
    return date

def split_date_components(date):
    return date.strftime('%B %d, %Y'), date.strftime('%I:%M:%S %p')

def set_font_color(run, hex_color):
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)

def create_summary_report(snapshots, domain, logo_path, logo_height_percent, title_color_hex, heading_color_hex, filename='summary_report.docx'):
    data = [(i+1, *split_date_components(extract_date_from_link(snap))) for i, snap in enumerate(snapshots)]
    df = pd.DataFrame(data, columns=['Capture', 'Date', 'Time'])

    summary = {
        'Total Captures': len(snapshots),
        'First Capture': {
            'Date': df['Date'].iloc[0],
            'Time': df['Time'].iloc[0]
        },
        'Last Capture': {
            'Date': df['Date'].iloc[-1],
            'Time': df['Time'].iloc[-1]
        }
    }

    doc = Document()
    
    # Add the logo above the title
    if logo_path:
        from PIL import Image
        logo = Image.open(logo_path)
        logo_width, logo_height = logo.size
        logo_height = int(logo_height * int(logo_height_percent) / 100)
        logo_width = int(logo_width * logo_height / logo.size[1])
        doc.add_picture(logo_path, width=Inches(logo_width / 96), height=Inches(logo_height / 96))

    title_run = doc.add_heading(f'Snapshot Summary Report for {domain}', 0).runs[0]
    set_font_color(title_run, title_color_hex)

    heading_1_run = doc.add_heading('Introduction', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph('This report provides a summary of the snapshots captured from the specified URL. The snapshots represent historical captures of the website, which are the first step in our process of identifying and fixing broken links. By analyzing these snapshots, we can gain insights into the changes and updates made to the website over time.')

    heading_1_run = doc.add_heading('Methodology', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph('The data was collected to provide a timeline of the website\'s state at various points in time. These snapshots help us understand how the website has evolved and where issues such as broken links may have arisen.')

    heading_1_run = doc.add_heading('Summary of Captures', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph(f"Total Captures: {summary['Total Captures']}")
    
    heading_2_run = doc.add_heading('First Capture', level=2).runs[0]
    set_font_color(heading_2_run, heading_color_hex)
    doc.add_paragraph(f"Date: {summary['First Capture']['Date']}")
    doc.add_paragraph(f"Time: {summary['First Capture']['Time']}")
    
    heading_2_run = doc.add_heading('Last Capture', level=2).runs[0]
    set_font_color(heading_2_run, heading_color_hex)
    doc.add_paragraph(f"Date: {summary['Last Capture']['Date']}")
    doc.add_paragraph(f"Time: {summary['Last Capture']['Time']}")

    doc.save(filename)

def create_detailed_analysis_report(snapshots, domain, logo_path, logo_height_percent, title_color_hex, heading_color_hex, column_color_hex, filename='detailed_analysis_report.docx'):
    data = [(i+1, *split_date_components(extract_date_from_link(snap))) for i, snap in enumerate(snapshots)]
    df = pd.DataFrame(data, columns=['Capture', 'Date', 'Time'])

    doc = Document()
    
    # Add the logo above the title
    if logo_path:
        from PIL import Image
        logo = Image.open(logo_path)
        logo_width, logo_height = logo.size
        logo_height = int(logo_height * int(logo_height_percent) / 100)
        logo_width = int(logo_width * logo_height / logo.size[1])
        doc.add_picture(logo_path, width=Inches(logo_width / 96), height=Inches(logo_height / 96))

    title_run = doc.add_heading(f'Detailed Snapshot Analysis for {domain}', 0).runs[0]
    set_font_color(title_run, title_color_hex)

    heading_1_run = doc.add_heading('Introduction', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph('This document provides a detailed analysis of the snapshots captured from the specified URL. Each snapshot represents a historical state of the website. This detailed analysis is a crucial step in our process of identifying and fixing broken links across the website by examining each snapshot and checking all the files and links within them.')

    heading_1_run = doc.add_heading('Detailed Analysis', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)
    doc.add_paragraph(f"Total Captures: {len(snapshots)}")
    doc.add_paragraph('The table below lists the snapshots in ascending order, organized by date and time.')

    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Capture'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Time'

    for cell in hdr_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(12)
        shading_elm_1 = OxmlElement("w:shd")
        shading_elm_1.set(qn("w:fill"), column_color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm_1)

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Capture'])
        row_cells[1].text = row['Date']
        row_cells[2].text = row['Time']

    # Add borders to the table
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    heading_1_run = doc.add_heading('Next Steps', level=1).runs[0]
    set_font_color(heading_1_run, heading_color_hex)

    def add_bold_paragraph(doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        return p

    add_bold_paragraph(doc, '1. Examination of Snapshots:')
    doc.add_paragraph('Each snapshot will be analyzed to check for broken links. This includes going through every file and link within each historical capture.')

    add_bold_paragraph(doc, '2. Identification and Fixing of Broken Links:')
    doc.add_paragraph('Any broken links identified during the analysis will be fixed.')

    add_bold_paragraph(doc, '3. Follow-Up Reports:')
    doc.add_paragraph('Additional reports will be provided throughout the process to document the progress and any further findings.')

    doc.save(filename)

def browse_file(variable):
    file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])
    variable.set(file_path)

def generate_reports():
    url = url_var.get()
    if not url.startswith("http://") and not url.startswith("https://"):
        url = "http://" + url
    domain = url.split('//')[1].split('/')[0]
    snapshots = get_snapshots(url)
    logo_path = logo_path_var.get()
    logo_height_percent = logo_height_percent_var.get()
    title_color_hex = title_color_hex_var.get().lstrip('#')
    heading_color_hex = heading_color_hex_var.get().lstrip('#')
    column_color_hex = column_color_hex_var.get().lstrip('#')
    if snapshots:
        create_summary_report(snapshots, domain, logo_path, logo_height_percent, title_color_hex, heading_color_hex)
        create_detailed_analysis_report(snapshots, domain, logo_path, logo_height_percent, title_color_hex, heading_color_hex, column_color_hex)
        messagebox.showinfo("Success", "Reports saved to summary_report.docx and detailed_analysis_report.docx")
    else:
        messagebox.showerror("Error", "No snapshots found or there was an error.")

if __name__ == '__main__':
    root = Tk()
    root.title("Snapshot Analysis Report Generator")
    root.geometry("600x450")

    url_var = StringVar()
    logo_path_var = StringVar()
    logo_height_percent_var = StringVar(value="50")
    title_color_hex_var = StringVar(value="000000")
    heading_color_hex_var = StringVar(value="000000")
    column_color_hex_var = StringVar(value="FFFFFF")

    Label(root, text="Enter URL:").grid(row=0, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=url_var, width=50).grid(row=0, column=1, padx=10, pady=10, sticky='w')

    Label(root, text="Logo File:").grid(row=1, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=logo_path_var, width=50).grid(row=1, column=1, padx=10, pady=10, sticky='w')
    Button(root, text="Browse", command=lambda: browse_file(logo_path_var)).grid(row=1, column=2, padx=10, pady=10, sticky='w')

    Label(root, text="Logo Height (%):").grid(row=2, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=logo_height_percent_var, width=10).grid(row=2, column=1, padx=10, pady=10, sticky='w')

    Label(root, text="Title Font Color (Hex):").grid(row=3, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=title_color_hex_var, width=10).grid(row=3, column=1, padx=10, pady=10, sticky='w')

    Label(root, text="Heading Font Color (Hex):").grid(row=4, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=heading_color_hex_var, width=10).grid(row=4, column=1, padx=10, pady=10, sticky='w')

    Label(root, text="Column Shading Color (Hex):").grid(row=5, column=0, padx=10, pady=10, sticky='w')
    Entry(root, textvariable=column_color_hex_var, width=10).grid(row=5, column=1, padx=10, pady=10, sticky='w')

    Button(root, text="Generate Reports", command=generate_reports).grid(row=6, columnspan=3, padx=10, pady=20)

    root.mainloop()
