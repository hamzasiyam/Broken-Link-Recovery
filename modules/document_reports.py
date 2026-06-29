"""Word document report generation for snapshot analysis workflows.

This module creates summary and detailed analysis reports using python-docx.
It contains the document-specific formatting logic so GUI code can simply pass
snapshot data, colors, and optional logo settings.
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from modules.paths import PROCESSED_DIR
from modules.wayback import extract_date_from_link, split_date_components


def create_summary_report(
    snapshots: list[str],
    domain: str,
    logo_path: str,
    logo_height_percent: str,
    title_color_hex: str,
    heading_color_hex: str,
    output_dir: str | Path = PROCESSED_DIR,
) -> Path:
    """Create a Word summary report for a domain's snapshots.

    Args:
        snapshots: List of Wayback snapshot URL strings.
        domain: Domain string used in the report title and filename.
        logo_path: Optional path string to an image file.
        logo_height_percent: String percentage used to scale the logo height.
        title_color_hex: Hex color string for the report title.
        heading_color_hex: Hex color string for section headings.
        output_dir: Directory path where the document should be saved.

    Returns:
        ``Path`` to the generated ``.docx`` file.

    Raises:
        ValueError: If ``snapshots`` is empty or a provided color is invalid.
    """
    # Stop early if there is no data to summarize.
    if not snapshots:
        raise ValueError("No snapshots were provided.")

    # Build the destination path and ensure the output directory exists.
    output_path = Path(output_dir) / f"{domain}_summary_report.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Parse all capture dates so the first and last captures can be calculated
    # even if the incoming list is not perfectly sorted.
    dates = [extract_date_from_link(snapshot) for snapshot in snapshots]
    first_date = min(dates)
    last_date = max(dates)
    first_capture_date, first_capture_time = split_date_components(first_date)
    last_capture_date, last_capture_time = split_date_components(last_date)

    # Start a new Word document and optionally insert the user-selected logo.
    document = Document()
    _add_logo(document, logo_path, logo_height_percent)

    # Add the title first and apply the user-selected title color.
    title_run = document.add_heading(f"Snapshot Summary Report for {domain}", 0).runs[0]
    set_font_color(title_run, title_color_hex)

    # Each heading helper call creates a section and applies consistent color.
    _add_heading(document, "Introduction", 1, heading_color_hex)
    document.add_paragraph(
        "This report provides a summary of the snapshots captured from the specified URL. "
        "The snapshots represent historical captures of the website, which are the first "
        "step in the process of identifying and fixing broken links. By analyzing these "
        "snapshots, the report shows how the website changed over time."
    )

    _add_heading(document, "Methodology", 1, heading_color_hex)
    document.add_paragraph(
        "The data was collected to provide a timeline of the website's state at various "
        "points in time. These snapshots help identify where broken links may have been "
        "introduced or removed."
    )

    _add_heading(document, "Summary of Captures", 1, heading_color_hex)
    document.add_paragraph(f"Total Captures: {len(snapshots)}")

    _add_heading(document, "First Capture", 2, heading_color_hex)
    document.add_paragraph(f"Date: {first_capture_date}")
    document.add_paragraph(f"Time: {first_capture_time}")

    _add_heading(document, "Last Capture", 2, heading_color_hex)
    document.add_paragraph(f"Date: {last_capture_date}")
    document.add_paragraph(f"Time: {last_capture_time}")

    # Save the completed document and return the path for GUI success messages.
    document.save(output_path)
    return output_path


def create_detailed_analysis_report(
    snapshots: list[str],
    domain: str,
    logo_path: str,
    logo_height_percent: str,
    title_color_hex: str,
    heading_color_hex: str,
    column_color_hex: str,
    output_dir: str | Path = PROCESSED_DIR,
) -> Path:
    """Create a detailed Word report listing every snapshot capture.

    Args:
        snapshots: List of Wayback snapshot URL strings.
        domain: Domain string used in the report title and filename.
        logo_path: Optional path string to an image file.
        logo_height_percent: String percentage used to scale the logo height.
        title_color_hex: Hex color string for the report title.
        heading_color_hex: Hex color string for section headings.
        column_color_hex: Hex color string used as the table header shading.
        output_dir: Directory path where the document should be saved.

    Returns:
        ``Path`` to the generated ``.docx`` file.

    Raises:
        ValueError: If ``snapshots`` is empty or a provided color is invalid.
    """
    # Stop early if there is no detailed data to list.
    if not snapshots:
        raise ValueError("No snapshots were provided.")

    # Build the destination path and ensure the report folder exists.
    output_path = Path(output_dir) / f"{domain}_detailed_analysis_report.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Convert snapshots into rows with a capture number, date, and time.
    data = [
        (index + 1, *split_date_components(extract_date_from_link(snapshot)))
        for index, snapshot in enumerate(snapshots)
    ]
    df = pd.DataFrame(data, columns=["Capture", "Date", "Time"])

    # Create the document and add optional branding before the title.
    document = Document()
    _add_logo(document, logo_path, logo_height_percent)

    title_run = document.add_heading(f"Detailed Snapshot Analysis for {domain}", 0).runs[0]
    set_font_color(title_run, title_color_hex)

    _add_heading(document, "Introduction", 1, heading_color_hex)
    document.add_paragraph(
        "This document provides a detailed analysis of the snapshots captured from the "
        "specified URL. Each snapshot represents a historical state of the website and "
        "helps guide broken-link analysis."
    )

    _add_heading(document, "Detailed Analysis", 1, heading_color_hex)
    document.add_paragraph(f"Total Captures: {len(snapshots)}")
    document.add_paragraph(
        "The table below lists the snapshots in ascending order, organized by date and time."
    )

    # Create a three-column table for all capture rows.
    table = document.add_table(rows=1, cols=3)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Capture"
    header_cells[1].text = "Date"
    header_cells[2].text = "Time"

    # Style the table header so it reads like a report table instead of raw data.
    for cell in header_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(12)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), clean_hex_color(column_color_hex))
        cell._element.get_or_add_tcPr().append(shading)

    # Append one Word table row for each snapshot in the pandas DataFrame.
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Capture"])
        row_cells[1].text = row["Date"]
        row_cells[2].text = row["Time"]

    _add_table_borders(table)

    # Add next-step guidance so generated reports are useful as handoff
    # documents for broken-link recovery work.
    _add_heading(document, "Next Steps", 1, heading_color_hex)
    _add_bold_paragraph(document, "1. Examination of Snapshots:")
    document.add_paragraph(
        "Each snapshot can be analyzed to check for broken links by reviewing the files "
        "and links in that historical capture."
    )

    _add_bold_paragraph(document, "2. Identification and Fixing of Broken Links:")
    document.add_paragraph("Broken links identified during analysis can be corrected.")

    _add_bold_paragraph(document, "3. Follow-Up Reports:")
    document.add_paragraph(
        "Additional reports can document progress and any findings from the repair process."
    )

    # Save the completed document and return the path for GUI success messages.
    document.save(output_path)
    return output_path


def set_font_color(run, hex_color: str) -> None:
    """Apply a hex color to a python-docx text run.

    Args:
        run: python-docx run object whose font color should be changed.
        hex_color: Six-character hex color string, optionally prefixed with
            ``#``.

    Returns:
        None. The run is modified in place.
    """
    # Validate and normalize the color before converting it into RGB integers.
    clean_color = clean_hex_color(hex_color)
    red, green, blue = tuple(int(clean_color[index : index + 2], 16) for index in (0, 2, 4))
    run.font.color.rgb = RGBColor(red, green, blue)


def clean_hex_color(hex_color: str) -> str:
    """Normalize and validate a six-character hex color.

    Args:
        hex_color: Color string such as ``"161747"`` or ``"#161747"``.

    Returns:
        Six-character hex color string without a leading ``#``.

    Raises:
        ValueError: If the color is not six valid hexadecimal characters.
    """
    # Default blank values to black so optional fields still produce a valid
    # report.
    clean_color = (hex_color or "000000").strip().lstrip("#")

    # If the color is not six characters, Word cannot use it as an RGB value.
    if len(clean_color) != 6:
        raise ValueError(f"Invalid hex color: {hex_color}")

    # Convert to int only to validate that every character is hexadecimal.
    int(clean_color, 16)
    return clean_color


def _add_logo(document: Document, logo_path: str, logo_height_percent: str) -> None:
    """Add an optional logo image to the top of a Word document.

    Args:
        document: python-docx ``Document`` object to modify.
        logo_path: Path string to the logo image. Empty strings are ignored.
        logo_height_percent: String percentage used to scale the original image.

    Returns:
        None. The document is modified in place if a valid logo exists.
    """
    # If the user did not choose a logo, leave the document unbranded.
    if not logo_path:
        return

    # If the stored logo path no longer exists, skip the logo instead of failing
    # the entire report generation.
    logo_file = Path(logo_path)
    if not logo_file.exists():
        return

    # Pillow is imported lazily because only report generation with a logo needs
    # image dimension inspection.
    from PIL import Image

    # Preserve the logo's aspect ratio while scaling its height by the requested
    # percentage.
    logo = Image.open(logo_file)
    logo_width, logo_height = logo.size
    height_percent = int(logo_height_percent or 50)
    scaled_height = int(logo_height * height_percent / 100)
    scaled_width = int(logo_width * scaled_height / logo_height)
    document.add_picture(
        str(logo_file),
        width=Inches(scaled_width / 96),
        height=Inches(scaled_height / 96),
    )


def _add_heading(document: Document, text: str, level: int, color_hex: str):
    """Add a colored heading to a Word document.

    Args:
        document: python-docx ``Document`` object to modify.
        text: Heading text string.
        level: Word heading level integer.
        color_hex: Hex color string for the heading run.

    Returns:
        The python-docx run object created for the heading.
    """
    # Add the heading, then style the first run because python-docx exposes
    # color formatting on runs rather than heading objects.
    heading_run = document.add_heading(text, level=level).runs[0]
    set_font_color(heading_run, color_hex)
    return heading_run


def _add_bold_paragraph(document: Document, text: str):
    """Add a paragraph containing one bold run.

    Args:
        document: python-docx ``Document`` object to modify.
        text: Paragraph text string.

    Returns:
        The python-docx paragraph object that was added.
    """
    # Create a paragraph and make the only run bold for numbered next-step
    # labels.
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    return paragraph


def _add_table_borders(table) -> None:
    """Apply simple borders to every cell in a Word table.

    Args:
        table: python-docx table object to modify.

    Returns:
        None. The table XML is modified in place.
    """
    # python-docx does not expose border styling directly, so update the
    # underlying table cell XML.
    for cell in table._tbl.iter_tcs():
        borders = OxmlElement("w:tcBorders")

        # Add the same border style to every side of every table cell.
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            borders.append(border)
        cell.tcPr.append(borders)
